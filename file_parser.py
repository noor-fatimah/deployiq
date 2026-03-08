# file_parser.py — DeployIQ
# Converts any uploaded file (CSV, Excel, PDF, Word, TXT, JSON) into a
# pandas DataFrame so the evaluator can work on it unchanged.

import io
import os
import json
import pandas as pd


SUPPORTED_EXTENSIONS = {
    ".csv",
    ".tsv",
    ".xlsx", ".xls", ".xlsm",
    ".pdf",
    ".docx", ".doc",
    ".txt",
    ".json",
}


def _ext(filename: str) -> str:
    return os.path.splitext(filename.lower())[1]


# ── CSV / TSV ──────────────────────────────────────────────────────────────────
def _parse_csv(path: str) -> pd.DataFrame:
    sep = "\t" if path.endswith(".tsv") else ","
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            return pd.read_csv(path, encoding=enc, sep=sep)
        except Exception:
            continue
    raise ValueError("Could not parse CSV/TSV file.")


# ── Excel ─────────────────────────────────────────────────────────────────────
def _parse_excel(path: str) -> pd.DataFrame:
    try:
        import openpyxl  # noqa — ensures xlsx support
    except ImportError:
        pass
    df = pd.read_excel(path, engine="openpyxl" if path.endswith(("xlsx", "xlsm")) else "xlrd")
    return df


# ── PDF ───────────────────────────────────────────────────────────────────────
def _parse_pdf(path: str) -> pd.DataFrame:
    """Extract tables from PDF; falls back to text extraction if no tables found."""
    # Try tabula first (best for structured tables)
    try:
        import tabula
        dfs = tabula.read_pdf(path, pages="all", multiple_tables=True, silent=True)
        # Merge all tables, keep only the largest one
        if dfs:
            dfs = [d for d in dfs if not d.empty and len(d.columns) >= 2]
        if dfs:
            df = max(dfs, key=lambda d: len(d))
            df.columns = [str(c).strip() for c in df.columns]
            return df
    except Exception:
        pass

    # Fallback: pdfplumber text → try to parse as CSV-like
    try:
        import pdfplumber
        rows = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    parts = [p.strip() for p in line.replace("\t", ",").split(",")]
                    if len(parts) >= 2:
                        rows.append(parts)
        if rows:
            # Use first row as header if it looks like a header
            header = rows[0]
            data   = rows[1:]
            if data:
                return pd.DataFrame(data, columns=header[:len(data[0])])
    except Exception:
        pass

    raise ValueError(
        "Could not extract a data table from the PDF. "
        "Please ensure the PDF contains a table with actual/predicted columns."
    )


# ── Word (.docx) ───────────────────────────────────────────────────────────────
def _parse_docx(path: str) -> pd.DataFrame:
    try:
        from docx import Document
        doc = Document(path)

        # Try tables in the document first
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if len(rows) >= 2 and len(rows[0]) >= 2:
                df = pd.DataFrame(rows[1:], columns=rows[0])
                if not df.empty:
                    return df

        # Fallback: paragraph text → CSV-like
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        rows = []
        for line in lines:
            parts = [x.strip() for x in line.replace("\t", ",").split(",")]
            if len(parts) >= 2:
                rows.append(parts)
        if rows:
            header = rows[0]
            data   = rows[1:]
            if data:
                return pd.DataFrame(data, columns=header[:len(data[0])])
    except ImportError:
        raise ValueError("python-docx is required to parse Word files. Add 'python-docx' to requirements.txt.")
    raise ValueError("Could not extract a data table from the Word document.")


# ── JSON ───────────────────────────────────────────────────────────────────────
def _parse_json(path: str) -> pd.DataFrame:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Support: list of dicts, dict of lists, or nested {"data": [...]}
    if isinstance(data, list):
        return pd.DataFrame(data)
    if isinstance(data, dict):
        # Try common wrappers
        for key in ("data", "records", "rows", "results", "items"):
            if key in data and isinstance(data[key], list):
                return pd.DataFrame(data[key])
        # Assume dict-of-lists / dict-of-values
        return pd.DataFrame(data)
    raise ValueError("JSON format not recognised. Expected a list of records or a dict of columns.")


# ── TXT ────────────────────────────────────────────────────────────────────────
def _parse_txt(path: str) -> pd.DataFrame:
    # Try tab-separated, then comma-separated, then space-separated
    for sep in ["\t", ",", " "]:
        try:
            df = pd.read_csv(path, sep=sep, engine="python")
            if len(df.columns) >= 2:
                return df
        except Exception:
            continue
    raise ValueError("Could not parse TXT file as tabular data.")


# ── Public API ─────────────────────────────────────────────────────────────────
def parse_file(path: str, filename: str) -> tuple[pd.DataFrame, str]:
    """
    Parse any supported file and return (DataFrame, detected_file_type).
    Raises ValueError with a user-friendly message on failure.
    """
    ext = _ext(filename)

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Accepted formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext in (".csv", ".tsv"):
        return _parse_csv(path), "csv"
    elif ext in (".xlsx", ".xls", ".xlsm"):
        return _parse_excel(path), "excel"
    elif ext == ".pdf":
        return _parse_pdf(path), "pdf"
    elif ext in (".docx", ".doc"):
        return _parse_docx(path), "word"
    elif ext == ".json":
        return _parse_json(path), "json"
    elif ext == ".txt":
        return _parse_txt(path), "txt"
    else:
        raise ValueError(f"No parser available for '{ext}'.")
